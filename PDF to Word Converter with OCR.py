import tkinter as tk
from tkinter import filedialog, messagebox
import os
from pdf2image import convert_from_path
import pytesseract
from docx import Document

# Fungsi untuk memilih file PDF
def select_pdf():
    file_path = filedialog.askopenfilename(
        filetypes=[("PDF Files", "*.pdf")], 
        title="Pilih File PDF"
    )
    if file_path:
        pdf_path_var.set(file_path)

# Fungsi untuk memilih direktori penyimpanan
def select_output_dir():
    dir_path = filedialog.askdirectory(title="Pilih Direktori Output")
    if dir_path:
        output_dir_var.set(dir_path)

# Fungsi untuk melakukan konversi
def convert_pdf():
    # Konfigurasi path Tesseract
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Ganti dengan lokasi tesseract.exe Anda
    
    pdf_path = pdf_path_var.get()
    output_dir = output_dir_var.get()
    use_ocr = ocr_var.get()

    if not pdf_path or not os.path.isfile(pdf_path):
        messagebox.showerror("Error", "Pilih file PDF yang valid.")
        return
    if not output_dir or not os.path.isdir(output_dir):
        messagebox.showerror("Error", "Pilih direktori output yang valid.")
        return

    try:
        # Nama file Word hasil konversi
        output_word_path = os.path.join(output_dir, "Hasil_Konversi.docx")
        document = Document()

        if use_ocr:
            # Konversi dengan OCR
            images = convert_from_path(pdf_path, dpi=300)
            for i, img in enumerate(images):
                text = pytesseract.image_to_string(img, lang="ind")  # Ganti "ind" jika ingin bahasa lain
                document.add_paragraph(f"Halaman {i + 1}")
                document.add_paragraph(text)
        else:
            # Konversi tanpa OCR (hanya file PDF dengan teks yang dapat diekstrak langsung)
            from PyPDF2 import PdfReader
            reader = PdfReader(pdf_path)
            for page_num, page in enumerate(reader.pages):
                document.add_paragraph(f"Halaman {page_num + 1}")
                document.add_paragraph(page.extract_text())

        # Simpan file Word
        document.save(output_word_path)
        messagebox.showinfo("Sukses", f"File berhasil dikonversi ke: {output_word_path}")

    except Exception as e:
        messagebox.showerror("Error", f"Terjadi kesalahan: {str(e)}")

# GUI dengan Tkinter
root = tk.Tk()
root.title("Konversi PDF ke Word")
root.geometry("500x300")

# Variabel
pdf_path_var = tk.StringVar()
output_dir_var = tk.StringVar()
ocr_var = tk.BooleanVar()

# Layout
tk.Label(root, text="Pilih File PDF:").pack(anchor="w", padx=10, pady=5)
tk.Entry(root, textvariable=pdf_path_var, width=50).pack(anchor="w", padx=10)
tk.Button(root, text="Browse", command=select_pdf).pack(anchor="w", padx=10)

tk.Label(root, text="Pilih Direktori Output:").pack(anchor="w", padx=10, pady=5)
tk.Entry(root, textvariable=output_dir_var, width=50).pack(anchor="w", padx=10)
tk.Button(root, text="Browse", command=select_output_dir).pack(anchor="w", padx=10)

tk.Checkbutton(root, text="Gunakan OCR (untuk PDF berbasis gambar)", variable=ocr_var).pack(anchor="w", padx=10, pady=5)

tk.Button(root, text="Konversi", command=convert_pdf, bg="blue", fg="white").pack(pady=20)

# Main loop
root.mainloop()
